from docx import Document


def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        return "\n".join(paragraphs)
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ""